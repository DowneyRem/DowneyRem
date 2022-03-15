#!/usr/bin/python
# -*- coding: UTF-8 -*-
# import os
import docx
from docx.api import Document
from docx.oxml.ns import qn  # 设置字体
from docx.shared import Pt, Cm, Inches  # 设置大小
from docx.shared import RGBColor  # 设置颜色
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE  # 自定样式
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 对齐方式


def FirstLineIndex(key="四号"):
	dict = {
		"初号":3.04, "小初":2.54, "一号":1.84, "小一":1.68,
		"二号":1.55, "小二":1.26, "三号":1.14, "小三":1.05,
		"四号":1.00, "小四":0.84, "五号":0.74, "小五":0.63,
		"六号":0.53, "小六":0.46, "七号":0.39, "八号":0.35}
	indetCM = dict.get(key)
	return indetCM

def FontSize(key):
	dict = {
		"初号":42, "小初":36, "一号":26, "小一":24,
		"二号":22, "小二":18, "三号":16, "小三":15,
		"四号":14, "小四":12, "五号":10.5, "小五":9,
		"六号":7.5, "小六":6.5, "七号":5.5, "八号":5}
	FontPt = dict.get(key)
	return FontPt

def SetStyleName(StyleName):
	dict = {
		"正文":"Normal", "标题":"Title", "副标题":"subtitle", "无间隔":"No Spacing",
		"标题1":"Heading 1", "标题2":"Heading 2", "标题3":"Heading 3", "标题4":"Heading 4"}
	try:
		StyleName = dict.get(StyleName)
	except:
		StyleName = styles.add_style(StyleName, WD_STYLE_TYPE.PARAGRAPH)
	print(StyleName)
	return StyleName

def Alignment(align):
	dict = {
		"左":WD_PARAGRAPH_ALIGNMENT.LEFT, "居左":WD_PARAGRAPH_ALIGNMENT.LEFT, "左对齐":WD_PARAGRAPH_ALIGNMENT.LEFT,
		"中":WD_PARAGRAPH_ALIGNMENT.CENTER, "居中":WD_PARAGRAPH_ALIGNMENT.CENTER, "居中对齐":WD_PARAGRAPH_ALIGNMENT.CENTER,
		"右":WD_PARAGRAPH_ALIGNMENT.RIGHT, "居右":WD_PARAGRAPH_ALIGNMENT.RIGHT, "右对齐":WD_PARAGRAPH_ALIGNMENT.RIGHT,
		"两端":WD_PARAGRAPH_ALIGNMENT.JUSTIFY, "两端对齐":WD_PARAGRAPH_ALIGNMENT.JUSTIFY,"":WD_PARAGRAPH_ALIGNMENT.LEFT}
	align = dict.get(align)
	return align


# 新建文件
document = Document()



latent_styles = document.styles.latent_styles
print(len(latent_styles))
latent_style_names = [ls.name for ls in latent_styles]
# print(latent_style_names)
latent_styles.default_to_locked = False
latent_styles.default_to_unhide_when_used = False

latent_style = latent_styles.add_latent_style("NORMAL INDENT")
latent_style.hidden = False
latent_style.priority = 2
latent_style.quick_style = True


p2 = document.add_paragraph("文件已经写入")
p2.style = "Normal"



### 正文字体 Normal
document.styles["Normal"].font.name = "宋体"
document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # 中日韩汉字
document.styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), u"Times New Roman")  # 拉丁字母

document.styles["Normal"].font.size = Pt(12)  # 小四
document.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)  # 黑色

document.styles["Normal"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
document.styles["Normal"].paragraph_format.space_before = Pt(0)  # 设置段前 0 磅
document.styles["Normal"].paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
document.styles["Normal"].paragraph_format.line_spacing = 1  # 设置行间距为单倍行距
document.styles["Normal"].paragraph_format.left_indent = Cm(0)  # 设置左缩进
document.styles["Normal"].paragraph_format.right_indent = Cm(0)  # 设置右缩进


document.styles["Heading 1"].font.name = "黑体"
document.styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # 中日韩汉字
document.styles["Heading 1"]._element.rPr.rFonts.set(qn("w:ascii"), u"Times New Roman")  # 拉丁字母
document.styles["Heading 1"].font.size = Pt(24)  # 小四
document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)  # 黑色


def SetStyle(StyleName, font1, font2):
	# 样式名称，中文字体，英文字体，字号大小，疑似无法设置字体
	StyleName = SetStyleName(StyleName)
	# if font2 == "":
	# 	font2 = font1
	# document.styles[StyleName].font.name = font1
	# document.styles[StyleName]._element.rPr.rFonts.set(qn("w:eastAsia"), font1)  # 中日韩汉字
	# document.styles[StyleName]._element.rPr.rFonts.set(qn("w:ascii"), font2)  # 拉丁字母
	document.styles[StyleName].font.color.rgb = RGBColor(0, 0, 0)  # 默认黑色
	

def SetStyleFirstLineIndent(StyleName, SizeText, FirstLine=0):
	StyleName = SetStyleName(StyleName)
	IndentCm = FirstLineIndex(SizeText)
	document.styles[StyleName].paragraph_format.first_line_indent = Cm(FirstLine * IndentCm)  # 设置首行缩进厘米数


def SetStyleFontColor(StyleName, R=0,G=0,B=0):
	StyleName = SetStyleName(StyleName)
	document.styles[StyleName].font.color.rgb = RGBColor(R, G, B)  # 默认黑色
	
	
def SetStyleSpace(StyleName, before, after, line_spacing=1):
	# 样式名称，段前距，段后距，多倍行距或磅值
	StyleName = SetStyleName(StyleName)
	document.styles[StyleName].paragraph_format.space_before = Pt(before)  # 设置段前 0 磅
	document.styles[StyleName].paragraph_format.space_after = Pt(after)  # 设置段后 0 磅
	document.styles[StyleName].paragraph_format.line_spacing = line_spacing  # 设置行间距为单倍行距


def SetStyleIndent(StyleName, SizeText, left=0, right=0):
	# 样式名称，字体大小，左右缩进的字符数
	StyleName = SetStyleName(StyleName)
	IndentCm = 0.5 * FirstLineIndex(SizeText) #单文字缩进厘米数
	document.styles[StyleName].paragraph_format.left_indent = Cm(left * IndentCm)  # 设置左缩进
	document.styles[StyleName].paragraph_format.right_indent = Cm(right * IndentCm)  # 设置右缩进
	

def AddHeading(text, level, font1, font2, SizeText):
	#文本，标题等级，字体1，字体2，字体大小
	run = document.add_heading('', level).add_run(text)
	run.font.size = Pt(FontSize(SizeText))
	run.font.name = font1
	run._element.rPr.rFonts.set(qn("w:eastAsia"), font1)  # 中日韩汉字
	run._element.rPr.rFonts.set(qn("w:ascii"), font2)  # 拉丁字母


### SetStyle("标题1", "黑体", "Times New Roman")
### SetStyle("标题2", "黑体", "Times New Roman")
### SetStyle("标题3", "黑体", "Times New Roman")
### SetStyle("正文", "宋体", "Times New Roman")

SetStyleFontColor("标题", R=0,G=0,B=0)
SetStyleFontColor("标题1", R=0,G=0,B=0)
SetStyleFontColor("标题2", R=0,G=0,B=0)
SetStyleFontColor("标题3", R=0,G=0,B=0)
SetStyleFontColor("正文", R=0,G=0,B=0)

SetStyleSpace("标题1", 0, 0, line_spacing=1)
SetStyleSpace("标题2", 0, 0, line_spacing=1)
SetStyleSpace("标题3", 0, 0, line_spacing=1)


AddHeading("标题", 0, "黑体", "Times New Roman", "一号")
AddHeading("一级标题", 1, "黑体", "Times New Roman", "三号")
AddHeading("二级标题", 2, "黑体", "Times New Roman", "小三")
AddHeading("三级标题", 3, "黑体", "Times New Roman", "四号")


p1 = document.add_paragraph("python-docx 是用于创建可修改 微软 Word 的一个 python 库，")


# 段前插入
author = p1.insert_paragraph_before("见鬼的作者")
author_format = author.paragraph_format
author_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

run2 = p1.add_run("提供全套的 Word 操作，")
run2 = p1.add_run("是最常用的 Word 工具。")



# 添加分页
document.add_page_break()


# 保存文件
try:
	# document.save("test.docx")
	document.save("D:\\Users\\Administrator\\Desktop\\1.docx")
	print("—" * 30)
	print("【文件已写入】")
except:
	print("【文件被占用】")
# finally:
# 	pass
